#!/usr/bin/env python3
"""
conway_generator.py — Tech Bro Conway Workflow Generator v2
=============================================================
Converts documents, web domains, and single webpages into Conway DSL-compliant
JSON workflow files ready for Tech Bro's WorkflowOrchestrator to execute.

Three modes:

  doc-to-workflow      Read a .docx, .pdf, or .md/.txt file and emit one or
                       more Conway JSON workflow files from the described process.

  web-to-workflow      Crawl an entire domain (via sitemap first, BFS fallback),
                       deduplicate workflow pages, and emit one Conway JSON file
                       per distinct workflow found.

  webpage-to-workflow  Consume a single URL and emit one Conway JSON file per
                       workflow / process found on that page.

Usage
-----
  python conway_generator.py doc-to-workflow \\
      --file path/to/process.pdf \\
      [--api-key sk-ant-...] \\
      [--output-dir ./workflows]

  python conway_generator.py web-to-workflow \\
      --url https://corp.sap.com \\
      [--max-pages 200] \\
      [--concurrency 6] \\
      [--api-key sk-ant-...] \\
      [--output-dir ./workflows]

  python conway_generator.py webpage-to-workflow \\
      --url https://corp.sap.com/some/process-guide \\
      [--api-key sk-ant-...] \\
      [--output-dir ./workflows]

API key priority: --api-key flag → ANTHROPIC_API_KEY env var

Dependencies
------------
  pip install requests beautifulsoup4 anthropic
  pip install python-docx          # for .docx support
  pip install pypdf                # for .pdf support
  pip install pillow pytesseract   # for OCR fallback on scanned PDFs (optional)
"""

import argparse
import hashlib
import json
import os
import re
import sys
import time
import threading
import xml.etree.ElementTree as ET
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urljoin, urlparse
from urllib.robotparser import RobotFileParser
from typing import Optional
from collections import defaultdict

# ============================================================================
# SECTION 1: CONWAY DSL KNOWLEDGE BASE
# ============================================================================

CONWAY_STEP_TYPES = [
    "navigate", "wait_for_load", "find_element", "click", "fill_form",
    "extract_data", "wait_for_human", "conditional", "loop",
    "screenshot", "scroll", "wait", "check_element", "hover"
]

SYSTEM_PROMPT = """You are an expert browser automation engineer for Tech Bro's WorkflowOrchestrator.
You convert process descriptions into precise, executable Conway DSL JSON workflows.

════════════════════════════════════════════════════════
CONWAY DSL — COMPLETE SPECIFICATION
════════════════════════════════════════════════════════

TOP-LEVEL WORKFLOW OBJECT:
{
  "id": "workflow-<timestamp>",
  "name": "Human Readable Name",
  "description": "What this workflow accomplishes in one sentence",
  "tags": ["category", "platform-name"],
  "steps": [ ...step objects... ]
}

════════════════════════════════════════════════════════
STEP TYPE REFERENCE — ALL FIELDS, ALL TYPES
════════════════════════════════════════════════════════

navigate:
  {
    "id": "navigate_to_portal",
    "type": "navigate",
    "description": "Open the SAP Fiori launchpad",
    "url": "https://example.com/path",
    "waitUntil": "domcontentloaded",   // load | domcontentloaded | networkidle
    "timeout": 30000,
    "critical": true
  }

wait_for_load:
  {
    "id": "wait_for_dashboard",
    "type": "wait_for_load",
    "description": "Wait for main content to appear",
    "selector": ".main-content",       // omit to wait for networkidle
    "timeout": 15000
  }

find_element:
  {
    "id": "find_submit_btn",
    "type": "find_element",
    "description": "Locate the submit button",
    "selector": "[data-testid='submit-btn']",
    "description": "submit button",
    "timeout": 10000,
    "retries": 3,
    "retryDelay": 1000
  }

click:
  {
    "id": "click_create_order",
    "type": "click",
    "description": "Click Create Purchase Order",
    "selector": "button:has-text('Create PO')",
    "waitAfter": 2000,
    "timeout": 5000,
    "retries": 3,
    "retryDelay": 1000,
    "critical": true
  }

fill_form:
  {
    "id": "fill_vendor_form",
    "type": "fill_form",
    "description": "Enter vendor and amount details",
    "fields": {
      "#vendor-id": "{{extract_vendor.extracted.vendor_id.text}}",
      "#amount": "",
      "input[name='description']": ""
    },
    "submit": "button[type='submit']",
    "typingDelay": 50,
    "retries": 3,
    "retryDelay": 1000
  }

extract_data:
  {
    "id": "extract_confirmation",
    "type": "extract_data",
    "description": "Capture confirmation number and status",
    "selectors": {
      "confirmation_number": ".confirmation-id",
      "status": "[data-field='status']",
      "total": ".order-total"
    }
  }

wait_for_human:
  {
    "id": "human_login",
    "type": "wait_for_human",
    "description": "User must complete SSO login",
    "prompt": {
      "title": "Login Required",
      "message": "Please log in with your corporate SSO credentials, then click Continue",
      "inputType": "text"             // text | choice
    },
    "selector": "#sso-login-frame",   // element to spotlight (optional)
    "targetSelector": "#username",    // field to auto-fill with human response (optional)
    "autoFill": true,                 // auto-type response into targetSelector
    "spotlight": true,                // dim page and highlight selector
    "showArrow": false,               // show pointing arrow to selector
    "screenshot": true                // capture screenshot before showing prompt
  }

  For choice-type human interventions:
  {
    "id": "human_approval_decision",
    "type": "wait_for_human",
    "prompt": {
      "title": "Approval Required",
      "message": "Review the order details above. Do you want to submit?",
      "inputType": "choice",
      "choices": ["Submit Order", "Cancel", "Save as Draft"]
    },
    "screenshot": true
  }

conditional:
  {
    "id": "check_for_error",
    "type": "conditional",
    "description": "Branch if error banner is visible",
    "condition": "element_exists",    // element_exists | JS expression using {{stepId.prop}}
    "selector": ".error-banner",      // required when condition === "element_exists"
    "if_true": [
      {
        "id": "handle_error",
        "type": "wait_for_human",
        "prompt": { "title": "Error Detected", "message": "An error appeared. Resolve it then continue." }
      }
    ],
    "if_false": [
      {
        "id": "take_success_screenshot",
        "type": "screenshot",
        "description": "Confirm clean state",
        "tags": ["success"]
      }
    ]
  }

  Nested conditionals are fully supported inside if_true / if_false branches.

  JS condition example (references prior step results via template syntax):
  {
    "id": "check_total_valid",
    "type": "conditional",
    "condition": "{{extract_total.extracted.amount.text}} > 0",
    "if_true": [...],
    "if_false": [...]
  }

loop:
  {
    "id": "process_line_items",
    "type": "loop",
    "description": "Submit each line item in the order",
    "items": "{{extract_line_items.extracted.rows}}",
    "steps": [
      {
        "id": "fill_line_item",
        "type": "fill_form",
        "description": "Enter line item details",
        "fields": { "#item-description": "{{item.text}}" }
      },
      {
        "id": "click_add_item",
        "type": "click",
        "selector": "button:has-text('Add Item')"
      }
    ]
  }

screenshot:
  {
    "id": "screenshot_confirmation",
    "type": "screenshot",
    "description": "Capture order confirmation for audit trail",
    "fullPage": false,
    "tags": ["confirmation", "audit"],
    "showInUI": true,
    "critical": false
  }

scroll:
  {
    "id": "scroll_to_submit",
    "type": "scroll",
    "description": "Scroll submit button into view",
    "direction": "to_element",        // down | up | to_element
    "selector": "button[type='submit']",  // required if direction === "to_element"
    "pixels": 500                     // used when direction is down | up
  }

wait:
  {
    "id": "wait_for_processing",
    "type": "wait",
    "description": "Allow backend processing time",
    "seconds": 3
  }

check_element:
  {
    "id": "check_success_banner",
    "type": "check_element",
    "description": "Verify success notification appeared",
    "selector": ".success-notification",
    "waitFor": true
  }

hover:
  {
    "id": "hover_account_menu",
    "type": "hover",
    "description": "Reveal account dropdown",
    "selector": ".account-menu-trigger"
  }

════════════════════════════════════════════════════════
TEMPLATE SYNTAX — REFERENCE PRIOR STEP RESULTS
════════════════════════════════════════════════════════
  {{stepId.property}}          — top-level result property
  {{stepId.extracted.key.text}} — from extract_data results
  {{stepId.human_input.payload}} — from wait_for_human results
  {{stepId.clicked}}           — boolean from click result

════════════════════════════════════════════════════════
SELECTOR PRIORITY (use in this order)
════════════════════════════════════════════════════════
  1. [data-testid="submit-button"]
  2. #stable-unique-id
  3. .descriptive-semantic-class
  4. button:has-text("Submit Order")
  5. [aria-label="Close dialog"]
  6. input[type="email"]
  ALWAYS add selectorAlternatives array for click/fill_form/find_element:
  "selectorAlternatives": ["[class*='submit']", "button:has-text('Save')"]

════════════════════════════════════════════════════════
STEP METADATA — APPLIES TO ALL STEPS
════════════════════════════════════════════════════════
  "critical": true/false   — false means failure is logged but workflow continues
  "retries": 3             — retry count for interactive steps
  "retryDelay": 1000       — ms between retries
  "timeout": 10000         — ms before timeout for waits/finds

════════════════════════════════════════════════════════
WORKFLOW GENERATION RULES — FOLLOW ALL OF THESE
════════════════════════════════════════════════════════
  1.  ALWAYS start: navigate → wait_for_load
  2.  ALWAYS end complex workflows with a screenshot (tags: ["confirmation","audit"])
  3.  Use wait_for_human for: SSO/login, MFA, CAPTCHA, file uploads, manual approvals
  4.  Use screenshot BEFORE every irreversible action (submit, delete, approve)
  5.  Use check_element BEFORE every conditional that tests existence
  6.  Add waitAfter (1500-3000ms) to clicks that trigger navigation or modal opens
  7.  Use conditional to handle error states — always provide both if_true AND if_false
  8.  Use loop when content describes "for each", "repeat for all", "per item" patterns
  9.  Use DESCRIPTIVE step IDs: "fill_vendor_details" not "step_4"
  10. Mark cosmetic/non-blocking steps with "critical": false
  11. Add selectorAlternatives to every click, fill_form, and find_element step
  12. Nested conditionals ARE supported — use them for complex branching logic
  13. One workflow = one logical user goal. Don't bundle unrelated tasks.
  14. If the process mentions a specific named platform, use its real URL.

RESPOND ONLY WITH VALID JSON.
No markdown fences, no explanation text, no comments inside the JSON.
Single workflow → return one object.
Multiple distinct workflows → return a JSON array of objects."""

EXTRACTION_PROMPT = """Analyze the following content and identify ALL distinct workflows,
processes, procedures, or step-by-step instructions present.

For EACH one found, generate a complete, immediately-executable Conway DSL workflow JSON.

RULES FOR THIS EXTRACTION:
- A workflow is any sequence of user actions that achieves a discrete goal on a website or app.
- Ignore purely conceptual background text with no actionable steps.
- If the platform has a known URL, use it in navigate steps. Otherwise use https://[platform-url].
- Be specific: "Submit Purchase Order in SAP Ariba" is better than "Submit a form".
- Preserve specific field names, button labels, and menu paths from the source text.
- When the source mentions "click X then Y then Z", that's at least 3 steps.
- When the source mentions a form, capture every field name mentioned.

Source URL / Document: {source}

Content:
---
{content}
---

{extra_instruction}

RESPOND ONLY WITH VALID JSON (single workflow object OR JSON array of workflow objects).
No markdown. No explanation. Pure JSON only."""


# ============================================================================
# SECTION 2: CLAUDE API CLIENT WITH RETRY + CHUNKING
# ============================================================================

class ClaudeClient:
    MODEL = "claude-sonnet-4-6"
    CHUNK_SIZE = 10_000       # characters per chunk (safe for context window)
    CHUNK_OVERLAP = 800       # overlap between chunks to avoid splitting mid-workflow
    MAX_RETRIES = 5
    RETRY_BASE_DELAY = 2.0    # seconds, doubled each retry

    def __init__(self, api_key: str):
        self.api_key = api_key
        self._lock = threading.Lock()  # serialize API calls from thread pool

    def _call_api(self, content: str, source: str, extra: str = "") -> list[dict]:
        """Single Claude API call with exponential backoff retry."""
        try:
            import anthropic
        except ImportError:
            sys.exit("Install anthropic: pip install anthropic")

        client = anthropic.Anthropic(api_key=self.api_key)
        user_msg = EXTRACTION_PROMPT.format(
            source=source,
            content=content,
            extra_instruction=extra,
        )

        last_error = None
        for attempt in range(1, self.MAX_RETRIES + 1):
            try:
                with self._lock:
                    response = client.messages.create(
                        model=self.MODEL,
                        max_tokens=4096,
                        system=SYSTEM_PROMPT,
                        messages=[{"role": "user", "content": user_msg}],
                    )

                raw = response.content[0].text.strip()
                raw = re.sub(r"^```json\s*", "", raw, flags=re.MULTILINE)
                raw = re.sub(r"^```\s*", "", raw, flags=re.MULTILINE)
                raw = re.sub(r"\s*```$", "", raw, flags=re.MULTILINE)
                raw = raw.strip()

                parsed = json.loads(raw)
                if isinstance(parsed, dict):
                    parsed = [parsed]
                return parsed

            except json.JSONDecodeError as e:
                print(f"    [claude] JSON parse error (attempt {attempt}/{self.MAX_RETRIES}): {e}")
                # Try to extract any valid JSON objects from the response
                recovered = self._recover_json(raw)
                if recovered:
                    print(f"    [claude] Recovered {len(recovered)} workflow(s) from partial JSON")
                    return recovered
                last_error = e

            except Exception as e:
                err_str = str(e)
                status = None
                # Extract HTTP status if present
                m = re.search(r"status[_\s]?code[:\s]+(\d+)|(\d{3})", err_str, re.I)
                if m:
                    status = int(m.group(1) or m.group(2))

                is_retryable = status in (429, 500, 502, 503, 529) or \
                               "rate_limit" in err_str.lower() or \
                               "overloaded" in err_str.lower() or \
                               "timeout" in err_str.lower()

                if is_retryable and attempt < self.MAX_RETRIES:
                    delay = self.RETRY_BASE_DELAY * (2 ** (attempt - 1))
                    # Check for Retry-After header hint in error message
                    ra = re.search(r"retry.after[:\s]+(\d+)", err_str, re.I)
                    if ra:
                        delay = max(delay, int(ra.group(1)))
                    print(f"    [claude] Retryable error (attempt {attempt}/{self.MAX_RETRIES}), "
                          f"waiting {delay:.0f}s: {err_str[:80]}")
                    time.sleep(delay)
                    last_error = e
                else:
                    raise

        raise last_error or RuntimeError("Max retries exceeded")

    def _recover_json(self, text: str) -> list[dict]:
        """Best-effort extraction of JSON objects from malformed response."""
        results = []
        # Try to find complete JSON objects using brace matching
        depth = 0
        start = None
        for i, ch in enumerate(text):
            if ch == '{':
                if depth == 0:
                    start = i
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0 and start is not None:
                    try:
                        obj = json.loads(text[start:i+1])
                        if "steps" in obj:
                            results.append(obj)
                    except json.JSONDecodeError:
                        pass
                    start = None
        return results

    def generate_workflows(self, content: str, source: str,
                           extra: str = "") -> list[dict]:
        """
        Main entry point. Chunks large content, calls Claude per chunk,
        deduplicates results, and returns a flat list of workflow dicts.
        """
        chunks = self._chunk_content(content)
        print(f"    [claude] {len(content):,} chars → {len(chunks)} chunk(s)")

        all_workflows: list[dict] = []
        seen_fingerprints: set[str] = set()

        for i, chunk in enumerate(chunks):
            chunk_extra = extra
            if len(chunks) > 1:
                chunk_extra += (
                    f"\n\nNOTE: This is chunk {i+1} of {len(chunks)} from the same document. "
                    "Only extract workflows whose steps are substantially contained in this chunk. "
                    "Do not invent content not present here."
                )
            try:
                workflows = self._call_api(chunk, source, chunk_extra)
                for wf in workflows:
                    fp = self._workflow_fingerprint(wf)
                    if fp not in seen_fingerprints:
                        seen_fingerprints.add(fp)
                        all_workflows.append(wf)
                    else:
                        print(f"    [dedup] Skipped duplicate workflow: {wf.get('name','?')}")
            except Exception as e:
                print(f"    [error] Chunk {i+1} failed after all retries: {e}")

        return all_workflows

    def _chunk_content(self, content: str) -> list[str]:
        """
        Split content into overlapping chunks at natural boundaries
        (paragraph breaks, numbered steps) to avoid cutting mid-workflow.
        """
        if len(content) <= self.CHUNK_SIZE:
            return [content]

        # Split on paragraph boundaries
        paragraphs = re.split(r'\n\s*\n', content)
        chunks = []
        current = []
        current_len = 0

        for para in paragraphs:
            para_len = len(para)
            if current_len + para_len > self.CHUNK_SIZE and current:
                chunk_text = "\n\n".join(current)
                chunks.append(chunk_text)
                # Overlap: keep last N chars worth of paragraphs
                overlap_text = chunk_text[-self.CHUNK_OVERLAP:]
                current = [overlap_text]
                current_len = len(overlap_text)
            current.append(para)
            current_len += para_len + 2  # +2 for \n\n

        if current:
            chunks.append("\n\n".join(current))

        return chunks

    @staticmethod
    def _workflow_fingerprint(wf: dict) -> str:
        """
        Content-based fingerprint for deduplication.
        Based on normalized name + step count + first/last step types.
        """
        name = re.sub(r'\s+', ' ', (wf.get('name') or '').lower().strip())
        steps = wf.get('steps') or []
        step_sig = f"{len(steps)}"
        if steps:
            step_sig += f"|{steps[0].get('type','')}|{steps[-1].get('type','')}"
        raw = f"{name}::{step_sig}"
        return hashlib.sha256(raw.encode()).hexdigest()[:16]


# ============================================================================
# SECTION 3: WORKFLOW VALIDATION, FIXING, ENHANCEMENT
# ============================================================================

def validate_workflow(wf: dict) -> list[str]:
    """Return list of validation warnings."""
    warnings = []
    if not wf.get("id"):
        warnings.append("Missing 'id'")
    if not wf.get("name"):
        warnings.append("Missing 'name'")
    steps = wf.get("steps")
    if not isinstance(steps, list) or len(steps) == 0:
        warnings.append("Missing or empty 'steps' array")
        return warnings

    for i, step in enumerate(steps):
        sid = step.get("id", f"index_{i}")
        if not step.get("id"):
            warnings.append(f"Step {i}: missing 'id'")
        stype = step.get("type")
        if not stype:
            warnings.append(f"Step {sid}: missing 'type'")
        elif stype not in CONWAY_STEP_TYPES:
            warnings.append(f"Step {sid}: unknown type '{stype}'")

        # Validate conditional branches recursively
        if stype == "conditional":
            for branch_name in ("if_true", "if_false"):
                branch = step.get(branch_name, [])
                for j, bstep in enumerate(branch):
                    if not bstep.get("type"):
                        warnings.append(f"Step {sid}.{branch_name}[{j}]: missing 'type'")

        # Validate loop
        if stype == "loop":
            if not step.get("items"):
                warnings.append(f"Step {sid}: loop missing 'items'")
            if not step.get("steps"):
                warnings.append(f"Step {sid}: loop missing 'steps'")

    return warnings


def fix_workflow(wf: dict) -> dict:
    """Apply safe automatic fixes to produce a valid Conway workflow."""
    ts = int(time.time() * 1000)

    if not wf.get("id"):
        wf["id"] = f"workflow-{ts}"
    if not wf.get("name"):
        wf["name"] = "Unnamed Workflow"
    if not wf.get("description"):
        wf["description"] = ""
    if not isinstance(wf.get("tags"), list):
        wf["tags"] = []
    if not isinstance(wf.get("steps"), list):
        wf["steps"] = []

    wf["steps"] = [_fix_step(s) for s in wf["steps"]]
    return wf


def _fix_step(step: dict) -> dict:
    """Fix a single step and its nested branches/loops."""
    if not step.get("id"):
        step["id"] = f"step_{int(time.time()*1000) % 100000}"
    if not step.get("type"):
        step["type"] = "wait"

    stype = step["type"]

    # Retry defaults for interactive steps
    if stype in ("click", "fill_form", "find_element"):
        step.setdefault("retries", 3)
        step.setdefault("retryDelay", 1000)

    # Timeout defaults
    if stype in ("wait_for_load", "find_element", "check_element", "navigate"):
        step.setdefault("timeout", 10000)

    # selectorAlternatives for element-targeting steps
    if stype in ("click", "fill_form", "find_element") and step.get("selector"):
        if not step.get("selectorAlternatives"):
            step["selectorAlternatives"] = _generate_selector_alternatives(
                step["selector"]
            )

    # Recursively fix conditional branches
    if stype == "conditional":
        step["if_true"]  = [_fix_step(s) for s in step.get("if_true",  [])]
        step["if_false"] = [_fix_step(s) for s in step.get("if_false", [])]

    # Recursively fix loop steps
    if stype == "loop":
        step["steps"] = [_fix_step(s) for s in step.get("steps", [])]

    return step


def _generate_selector_alternatives(primary: str) -> list[str]:
    """Mirror the logic from AIOrchestrator.generateSelectorAlternatives."""
    alts = []
    # Text-based selector variants
    if ":has-text" in primary or ":text" in primary:
        m = re.search(r'["\'](.+?)["\']', primary)
        if m:
            text = m.group(1)
            alts += [
                f'button:contains("{text}")',
                f'a:contains("{text}")',
                f'[title*="{text}"]',
                f'[aria-label*="{text}"]',
            ]
    # Class selector
    if primary.startswith("."):
        cls = primary[1:]
        alts += [
            f'[class*="{cls}"]',
            f'*[class*="{cls.split("-")[0]}"]',
        ]
    # ID selector
    if primary.startswith("#"):
        id_ = primary[1:]
        alts += [
            f'[id="{id_}"]',
            f'[data-id="{id_}"]',
        ]
    # data-testid
    if "data-testid" in primary:
        m = re.search(r'data-testid=["\'](.+?)["\']', primary)
        if m:
            val = m.group(1)
            alts += [
                f'[data-test="{val}"]',
                f'[data-cy="{val}"]',
                f'[id*="{val}"]',
            ]
    return alts[:4]  # cap at 4 alternatives


# ============================================================================
# SECTION 4: MANIFEST
# ============================================================================

class Manifest:
    """
    Tracks every workflow file generated in a run.
    Written to {output_dir}/manifest.json on completion.
    """
    def __init__(self, output_dir: Path, mode: str, source: str):
        self.output_dir = output_dir
        self.mode = mode
        self.source = source
        self.started_at = datetime.now(timezone.utc).isoformat()
        self.entries: list[dict] = []
        self._lock = threading.Lock()

    def record(self, wf: dict, filepath: Path, source_url: str = "",
               chunk_index: Optional[int] = None):
        steps = wf.get("steps", [])
        entry = {
            "file": filepath.name,
            "workflow_id": wf.get("id", ""),
            "name": wf.get("name", ""),
            "description": wf.get("description", ""),
            "tags": wf.get("tags", []),
            "step_count": len(steps),
            "step_types": list({s.get("type") for s in steps if s.get("type")}),
            "source_url": source_url,
            "generated_at": datetime.now(timezone.utc).isoformat(),
        }
        if chunk_index is not None:
            entry["chunk_index"] = chunk_index
        with self._lock:
            self.entries.append(entry)

    def write(self):
        manifest = {
            "generator": "conway_generator.py v2",
            "mode": self.mode,
            "source": self.source,
            "started_at": self.started_at,
            "completed_at": datetime.now(timezone.utc).isoformat(),
            "total_workflows": len(self.entries),
            "output_dir": str(self.output_dir),
            "workflows": sorted(self.entries, key=lambda e: e["name"].lower()),
        }
        path = self.output_dir / "manifest.json"
        path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
        print(f"\n[manifest] Written: {path}  ({len(self.entries)} workflows)")
        return path


# ============================================================================
# SECTION 5: DOCUMENT READERS
# ============================================================================

def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        sys.exit("Install python-docx: pip install python-docx")

    doc = Document(str(path))
    parts = []

    # Headings and paragraphs with style context
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if "Heading" in style:
            parts.append(f"\n## {text}\n")
        elif "List" in style:
            parts.append(f"  - {text}")
        else:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        sys.exit("Install pypdf: pip install pypdf")

    reader = PdfReader(str(path))
    pages = []
    empty_pages = 0

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages.append(f"[Page {i+1}]\n{text}")
        else:
            empty_pages += 1

    if empty_pages > len(reader.pages) * 0.5:
        # More than half the pages are empty → likely a scanned PDF
        ocr_text = _ocr_pdf_fallback(path)
        if ocr_text:
            print(f"  [pdf] Used OCR fallback ({empty_pages} empty pages detected)")
            return ocr_text
        else:
            print(f"  [pdf] WARNING: {empty_pages}/{len(reader.pages)} pages appear to be scanned images.")
            print("        Install OCR support: pip install pillow pytesseract")
            print("        (also requires Tesseract: https://github.com/tesseract-ocr/tesseract)")

    return "\n\n".join(pages)


def _ocr_pdf_fallback(path: Path) -> str:
    """Attempt OCR on a scanned PDF using pytesseract + pillow."""
    try:
        import pytesseract
        from PIL import Image
        import io
    except ImportError:
        return ""

    try:
        from pypdf import PdfReader
        import subprocess

        # Try pdf2image if available
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(str(path), dpi=200)
            pages = []
            for i, img in enumerate(images):
                text = pytesseract.image_to_string(img)
                if text.strip():
                    pages.append(f"[Page {i+1} OCR]\n{text.strip()}")
            return "\n\n".join(pages)
        except ImportError:
            pass

        return ""
    except Exception as e:
        print(f"  [pdf] OCR failed: {e}")
        return ""


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def read_document(path: Path) -> str:
    suffix = path.suffix.lower()
    print(f"[doc] Reading {suffix}: {path.name}")
    if suffix == ".docx":
        return read_docx(path)
    elif suffix == ".pdf":
        return read_pdf(path)
    elif suffix in (".md", ".txt", ".rst", ".text"):
        return read_text(path)
    else:
        sys.exit(f"Unsupported file type: {suffix}. Supported: .docx .pdf .md .txt")


# ============================================================================
# SECTION 6: WEB UTILITIES — SESSION, ROBOTS, SITEMAP, CRAWL
# ============================================================================

def get_session():
    try:
        import requests
        s = requests.Session()
        s.headers.update({"User-Agent": "Mozilla/5.0 (compatible; ConwayBot/2.0)"})
        return s
    except ImportError:
        sys.exit("Install requests: pip install requests")


def fetch_page(session, url: str, timeout: int = 15) -> Optional[str]:
    try:
        resp = session.get(url, timeout=timeout)
        ct = resp.headers.get("content-type", "")
        if "text/html" not in ct and "text/plain" not in ct:
            return None
        return resp.text
    except Exception as e:
        print(f"    [fetch] {url}: {e}")
        return None


def fetch_xml(session, url: str, timeout: int = 10) -> Optional[str]:
    try:
        resp = session.get(url, timeout=timeout)
        if resp.status_code == 200:
            return resp.text
        return None
    except Exception:
        return None


def build_robots_parser(session, base_url: str) -> Optional[RobotFileParser]:
    """Fetch and parse robots.txt. Returns None if unavailable."""
    robots_url = urljoin(base_url, "/robots.txt")
    try:
        rp = RobotFileParser()
        rp.set_url(robots_url)
        xml = fetch_xml(session, robots_url)
        if xml:
            rp.parse(xml.splitlines())
            print(f"[robots] Loaded robots.txt from {robots_url}")
            return rp
    except Exception as e:
        print(f"[robots] Could not parse robots.txt: {e}")
    return None


def can_fetch(robots: Optional[RobotFileParser], url: str) -> bool:
    """Return True if robots.txt permits crawling this URL."""
    if robots is None:
        return True
    return robots.can_fetch("ConwayBot", url) or robots.can_fetch("*", url)


def discover_sitemap_urls(session, base_url: str) -> list[str]:
    """
    Try to find all URLs via sitemap.xml / sitemap_index.xml.
    Returns a flat list of all discovered URLs on the same domain.
    """
    candidates = [
        urljoin(base_url, "/sitemap.xml"),
        urljoin(base_url, "/sitemap_index.xml"),
        urljoin(base_url, "/sitemap-index.xml"),
        urljoin(base_url, "/sitemaps/sitemap.xml"),
    ]

    found_urls: list[str] = []
    processed_sitemaps: set[str] = set()
    sitemap_queue: list[str] = []

    # First check robots.txt for Sitemap: directives
    robots_txt = fetch_xml(session, urljoin(base_url, "/robots.txt"))
    if robots_txt:
        for line in robots_txt.splitlines():
            if line.lower().startswith("sitemap:"):
                sm_url = line.split(":", 1)[1].strip()
                if sm_url not in candidates:
                    candidates.insert(0, sm_url)

    for candidate in candidates:
        if candidate not in processed_sitemaps:
            sitemap_queue.append(candidate)

    def parse_sitemap(xml_text: str, source_url: str):
        """Parse sitemap XML, returns (page_urls, child_sitemap_urls)."""
        page_urls = []
        child_sitemaps = []
        try:
            # Strip namespaces for simpler parsing
            xml_clean = re.sub(r'\s+xmlns[^=]*="[^"]*"', '', xml_text)
            root = ET.fromstring(xml_clean)
            tag = root.tag.lower().split("}")[-1] if "}" in root.tag else root.tag.lower()

            if tag == "sitemapindex":
                for sitemap in root.iter():
                    t = sitemap.tag.split("}")[-1] if "}" in sitemap.tag else sitemap.tag
                    if t == "loc" and sitemap.text:
                        child_sitemaps.append(sitemap.text.strip())
            else:
                # Regular sitemap
                for url_elem in root.iter():
                    t = url_elem.tag.split("}")[-1] if "}" in url_elem.tag else url_elem.tag
                    if t == "loc" and url_elem.text:
                        page_urls.append(url_elem.text.strip())
        except ET.ParseError as e:
            print(f"    [sitemap] XML parse error at {source_url}: {e}")
        return page_urls, child_sitemaps

    while sitemap_queue:
        sm_url = sitemap_queue.pop(0)
        if sm_url in processed_sitemaps:
            continue
        processed_sitemaps.add(sm_url)

        xml = fetch_xml(session, sm_url)
        if not xml:
            continue

        print(f"  [sitemap] Parsing: {sm_url}")
        pages, children = parse_sitemap(xml, sm_url)

        # Filter to same domain
        domain_pages = [u for u in pages if same_domain(u, base_url)]
        found_urls.extend(domain_pages)

        for child in children:
            if child not in processed_sitemaps:
                sitemap_queue.append(child)

    if found_urls:
        print(f"[sitemap] Discovered {len(found_urls)} URLs via sitemap")
    else:
        print("[sitemap] No URLs found via sitemap — will use BFS crawl")

    return found_urls


def parse_html(html: str):
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        sys.exit("Install beautifulsoup4: pip install beautifulsoup4")
    return BeautifulSoup(html, "html.parser")


def extract_text(soup) -> str:
    """Extract clean, structured text from a BeautifulSoup page."""
    for tag in soup.find_all(["script", "style", "nav", "footer",
                               "header", "aside", "noscript", "svg"]):
        tag.decompose()

    main = (
        soup.find("main") or
        soup.find(attrs={"role": "main"}) or
        soup.find(id=re.compile(r"content|main|article|body", re.I)) or
        soup.find(class_=re.compile(r"content|main|article|body", re.I)) or
        soup.find("article") or
        soup.body or soup
    )

    lines = []
    for elem in main.find_all(
        ["h1","h2","h3","h4","h5","h6","p","li","td","th","dt","dd","pre","blockquote"]
    ):
        text = elem.get_text(separator=" ", strip=True)
        tag = elem.name
        if not text or len(text) < 10:
            continue
        if tag in ("h1","h2","h3"):
            lines.append(f"\n## {text}\n")
        elif tag in ("h4","h5","h6"):
            lines.append(f"\n### {text}\n")
        elif tag == "li":
            lines.append(f"  - {text}")
        else:
            lines.append(text)

    return "\n".join(lines)


def is_workflow_page(text: str, title: str = "") -> bool:
    """
    Heuristic: does this page likely describe a step-by-step process?
    Uses both pattern matching and a simple scoring system.
    """
    combined = (title + " " + text).lower()
    strong_patterns = [
        r'\bstep\s+\d+\b',
        r'\b\d+\.\s+[a-z]',
        r'\bhow\s+to\b',
        r'\bprocedure\b',
        r'\bstep[-\s]by[-\s]step\b',
        r'\binstructions?\b',
        r'\bworkflow\b',
        r'\bquick\s+start\b',
        r'\bwalkthrough\b',
        r'\btutorial\b',
        r'\bsubmit\s+a\b',
        r'\bcreate\s+a\b',
        r'\bconfigure\b',
        r'\bset\s+up\b',
        r'\bget\s+started\b',
    ]
    score = sum(1 for p in strong_patterns if re.search(p, combined))
    # Also score for numbered list density
    numbered = len(re.findall(r'^\s*\d+[\.\)]\s', text, re.MULTILINE))
    if numbered >= 3:
        score += 2
    return score >= 2


def same_domain(url: str, base: str) -> bool:
    return urlparse(url).netloc == urlparse(base).netloc


def normalize_url(href: str, current: str) -> Optional[str]:
    try:
        full = urljoin(current, href)
        p = urlparse(full)
        if p.scheme not in ("http", "https"):
            return None
        return p._replace(fragment="").geturl()
    except Exception:
        return None


def content_fingerprint(text: str) -> str:
    """Fingerprint page content for cross-page deduplication."""
    # Normalize whitespace, lowercase, hash first 3000 chars of normalized text
    normalized = re.sub(r'\s+', ' ', text.lower().strip())[:3000]
    return hashlib.sha256(normalized.encode()).hexdigest()[:16]


# ============================================================================
# SECTION 7: OUTPUT HELPERS
# ============================================================================

def slug(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")[:60]


def write_workflow(wf: dict, output_dir: Path, prefix: str = "") -> Path:
    wf = fix_workflow(wf)
    warnings = validate_workflow(wf)
    for w in warnings:
        print(f"    [warn] {w}")

    name_slug = slug(wf.get("name", "workflow"))
    base = f"{prefix}{name_slug}" if prefix else name_slug
    path = output_dir / f"{base}.json"

    counter = 1
    while path.exists():
        # Check if it's truly a different workflow (not the same one being re-written)
        try:
            existing = json.loads(path.read_text())
            if existing.get("id") == wf.get("id"):
                break  # Same workflow, overwrite
        except Exception:
            pass
        path = output_dir / f"{base}_{counter}.json"
        counter += 1

    output_dir.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(wf, indent=2), encoding="utf-8")
    print(f"    [out] {path.name}  ({len(wf['steps'])} steps)")
    return path


# ============================================================================
# SECTION 8: DOC-TO-WORKFLOW
# ============================================================================

def cmd_doc_to_workflow(args):
    api_key = _get_api_key(args)
    path = Path(args.file)
    if not path.exists():
        sys.exit(f"File not found: {path}")

    content = read_document(path)
    if not content.strip():
        sys.exit("Document appears empty after extraction.")

    print(f"[doc] Extracted {len(content):,} characters from '{path.name}'")

    out = Path(args.output_dir)
    manifest = Manifest(out, "doc-to-workflow", str(path.resolve()))
    client = ClaudeClient(api_key)

    extra = (
        f"This content is from a document named '{path.name}'. "
        "Extract EVERY distinct process, procedure, or workflow described — "
        "even if they seem minor. Preserve all specific field names, button labels, "
        "and menu paths mentioned in the source."
    )

    print(f"[generate] Sending to Claude...")
    workflows = client.generate_workflows(content, source=str(path), extra=extra)

    print(f"\n[doc] Generating {len(workflows)} workflow file(s)...")
    file_prefix = slug(path.stem) + "_"
    for wf in workflows:
        fp = write_workflow(wf, out, prefix=file_prefix)
        manifest.record(wf, fp, source_url=str(path))

    manifest.write()
    print(f"\n[done] {len(workflows)} workflow(s) written to {out}/")


# ============================================================================
# SECTION 9: WEB-TO-WORKFLOW
# ============================================================================

def cmd_web_to_workflow(args):
    api_key = _get_api_key(args)
    base_url = args.url.rstrip("/")
    max_pages = args.max_pages
    concurrency = args.concurrency
    out = Path(args.output_dir)

    session = get_session()
    manifest = Manifest(out, "web-to-workflow", base_url)
    client = ClaudeClient(api_key)

    # Step 1: Load robots.txt
    robots = build_robots_parser(session, base_url)
    crawl_delay = 0.5
    if robots:
        delay = robots.crawl_delay("ConwayBot") or robots.crawl_delay("*")
        if delay:
            crawl_delay = max(crawl_delay, float(delay))
            print(f"[robots] Crawl delay: {crawl_delay}s")

    # Step 2: Discover URLs via sitemap, fall back to BFS
    all_urls = discover_sitemap_urls(session, base_url)

    if all_urls:
        # Filter by robots.txt and same domain
        all_urls = [u for u in all_urls
                    if same_domain(u, base_url) and can_fetch(robots, u)]
        all_urls = list(dict.fromkeys(all_urls))  # deduplicate, preserve order
        print(f"[urls] {len(all_urls)} URLs after robots filter")
    else:
        # BFS crawl fallback
        print(f"[crawl] Starting BFS from {base_url} (max {max_pages} pages)")
        all_urls = _bfs_crawl(session, base_url, max_pages, robots, crawl_delay)

    # Limit to max_pages
    all_urls = all_urls[:max_pages]
    print(f"[crawl] Processing {len(all_urls)} URLs total")

    # Step 3: Filter to workflow pages (parallel fetch + heuristic)
    workflow_pages: list[tuple[str, str, str]] = []  # (url, text, title)
    page_content_fingerprints: set[str] = set()
    fetch_lock = threading.Lock()

    def fetch_and_classify(url: str) -> Optional[tuple]:
        if not can_fetch(robots, url):
            return None
        html = fetch_page(session, url)
        if not html:
            return None
        soup = parse_html(html)
        title = (soup.find("title") or soup.find("h1") or "")
        title_text = title.get_text(strip=True) if hasattr(title, 'get_text') else ""
        text = extract_text(soup)
        if not text.strip():
            return None
        fp = content_fingerprint(text)
        with fetch_lock:
            if fp in page_content_fingerprints:
                print(f"  [dedup] Skipping duplicate content: {url}")
                return None
            page_content_fingerprints.add(fp)
        if is_workflow_page(text, title_text):
            return (url, text, title_text)
        return None

    print(f"\n[filter] Fetching and classifying pages (concurrency={concurrency})...")
    with ThreadPoolExecutor(max_workers=concurrency) as pool:
        futures = {pool.submit(fetch_and_classify, url): url for url in all_urls}
        for i, future in enumerate(as_completed(futures), 1):
            url = futures[future]
            try:
                result = future.result()
                if result:
                    workflow_pages.append(result)
                    print(f"  [{i:>4}/{len(all_urls)}] ✓ workflow: {url}")
                else:
                    if i % 20 == 0:
                        print(f"  [{i:>4}/{len(all_urls)}] scanning...")
            except Exception as e:
                print(f"  [error] {url}: {e}")
            time.sleep(crawl_delay / concurrency)  # distributed rate respect

    print(f"\n[filter] {len(workflow_pages)} workflow pages found from {len(all_urls)} total\n")

    # Step 4: Generate Conway workflows per page
    total_written = 0
    domain_slug = slug(urlparse(base_url).netloc)

    # Deduplicate workflow-level too (same workflow documented on multiple pages)
    seen_workflow_fingerprints: set[str] = set()

    def generate_for_page(page_tuple: tuple) -> list[Path]:
        url, text, title = page_tuple
        extra = (
            f"This content is from: {url}\n"
            f"Page title: {title}\n"
            "The platform URL is known — use it directly in navigate steps. "
            "Extract every distinct workflow on this page. Be specific and complete."
        )
        try:
            workflows = client.generate_workflows(text, source=url, extra=extra)
            paths = []
            for wf in workflows:
                fp = ClaudeClient._workflow_fingerprint(wf)
                if fp in seen_workflow_fingerprints:
                    print(f"    [dedup] Duplicate workflow skipped: {wf.get('name','?')}")
                    continue
                seen_workflow_fingerprints.add(fp)
                p = write_workflow(wf, out, prefix=f"{domain_slug}_")
                manifest.record(wf, p, source_url=url)
                paths.append(p)
            return paths
        except Exception as e:
            print(f"    [error] Failed for {url}: {e}")
            return []

    print(f"[generate] Generating workflows ({concurrency} concurrent)...")
    with ThreadPoolExecutor(max_workers=max(1, concurrency // 2)) as pool:
        futures = {pool.submit(generate_for_page, pg): pg[0] for pg in workflow_pages}
        for future in as_completed(futures):
            url = futures[future]
            try:
                paths = future.result()
                total_written += len(paths)
                print(f"  [done] {url} → {len(paths)} workflow(s)")
            except Exception as e:
                print(f"  [error] {url}: {e}")

    manifest.write()
    print(f"\n[done] {total_written} workflow(s) from {len(workflow_pages)} pages → {out}/")


def _bfs_crawl(session, base_url: str, max_pages: int,
               robots: Optional[RobotFileParser], delay: float) -> list[str]:
    """BFS crawl fallback when no sitemap is available."""
    visited: set[str] = set()
    queue = [base_url]
    found: list[str] = []

    while queue and len(visited) < max_pages:
        url = queue.pop(0)
        if url in visited or not can_fetch(robots, url):
            continue
        visited.add(url)
        found.append(url)

        html = fetch_page(session, url)
        if not html:
            continue
        soup = parse_html(html)
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if any(href.startswith(x) for x in ("mailto:", "tel:", "javascript:", "#")):
                continue
            full = normalize_url(href, url)
            if full and same_domain(full, base_url) and full not in visited:
                queue.append(full)

        time.sleep(delay)

    print(f"[bfs] Crawled {len(found)} pages")
    return found


# ============================================================================
# SECTION 10: WEBPAGE-TO-WORKFLOW
# ============================================================================

def cmd_webpage_to_workflow(args):
    api_key = _get_api_key(args)
    url = args.url
    out = Path(args.output_dir)

    print(f"[page] Fetching {url}")
    session = get_session()
    html = fetch_page(session, url)
    if not html:
        sys.exit(f"Failed to fetch page: {url}")

    soup = parse_html(html)
    title = (soup.find("title") or soup.find("h1") or "")
    title_text = title.get_text(strip=True) if hasattr(title, 'get_text') else url
    text = extract_text(soup)

    if not text.strip():
        sys.exit("Page appears empty after text extraction.")

    print(f"[page] Extracted {len(text):,} characters | Title: {title_text}")

    if not is_workflow_page(text, title_text) and not args.force:
        print("[page] Warning: this page doesn't strongly appear to contain workflows.")
        print("       Use --force to generate anyway.")
        if not _confirm("Continue anyway? [y/N] "):
            sys.exit(0)

    manifest = Manifest(out, "webpage-to-workflow", url)
    client = ClaudeClient(api_key)

    extra = (
        f"Page title: {title_text}\n"
        f"Source URL: {url}\n"
        "Find and extract EVERY distinct workflow, process, or procedure on this page. "
        "If there are multiple unrelated workflows, return a JSON array. "
        "Preserve exact field names, button labels, and menu paths from the source."
    )

    print(f"[generate] Sending to Claude...")
    workflows = client.generate_workflows(text, source=url, extra=extra)

    print(f"\n[page] Generating {len(workflows)} workflow file(s)...")
    page_slug = slug(urlparse(url).path.strip("/") or urlparse(url).netloc)
    for wf in workflows:
        fp = write_workflow(wf, out, prefix=f"{page_slug}_")
        manifest.record(wf, fp, source_url=url)

    manifest.write()
    print(f"\n[done] {len(workflows)} workflow(s) written to {out}/")


# ============================================================================
# SECTION 11: SHARED UTILITIES
# ============================================================================

def _get_api_key(args) -> str:
    key = getattr(args, "api_key", None) or os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        sys.exit(
            "\nNo API key found.\n"
            "Set ANTHROPIC_API_KEY environment variable or pass --api-key <key>\n"
        )
    return key


def _confirm(prompt: str) -> bool:
    try:
        return input(prompt).strip().lower() in ("y", "yes")
    except (EOFError, KeyboardInterrupt):
        return False


# ============================================================================
# SECTION 12: CLI
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Conway Workflow Generator v2 — doc / web / webpage → Tech Bro JSON",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument(
        "--api-key", default=None,
        help="Anthropic API key (default: ANTHROPIC_API_KEY env var)"
    )
    parser.add_argument(
        "--output-dir", default="./workflows",
        help="Directory to write Conway JSON files (default: ./workflows)"
    )

    sub = parser.add_subparsers(dest="mode", required=True)

    # ── doc-to-workflow ──────────────────────────────────────────────────────
    dp = sub.add_parser(
        "doc-to-workflow",
        help="Read a document (.docx/.pdf/.md/.txt) and generate Conway workflow JSON(s)"
    )
    dp.add_argument(
        "--file", required=True,
        help="Path to document (.docx, .pdf, .md, .txt)"
    )

    # ── web-to-workflow ──────────────────────────────────────────────────────
    wp = sub.add_parser(
        "web-to-workflow",
        help="Crawl a domain (sitemap-first) and generate one Conway JSON per workflow found"
    )
    wp.add_argument("--url", required=True, help="Base URL to crawl")
    wp.add_argument(
        "--max-pages", type=int, default=200,
        help="Max pages to process (default: 200)"
    )
    wp.add_argument(
        "--concurrency", type=int, default=6,
        help="Parallel fetch/generate workers (default: 6)"
    )

    # ── webpage-to-workflow ──────────────────────────────────────────────────
    sp = sub.add_parser(
        "webpage-to-workflow",
        help="Consume a single URL and generate Conway workflow JSON(s)"
    )
    sp.add_argument("--url", required=True, help="URL of the page to process")
    sp.add_argument(
        "--force", action="store_true",
        help="Generate even if page doesn't appear to contain workflows"
    )

    args = parser.parse_args()

    print(f"\n{'═'*62}")
    print(f"  Conway Workflow Generator v2  |  mode: {args.mode}")
    print(f"  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"{'═'*62}\n")

    if args.mode == "doc-to-workflow":
        cmd_doc_to_workflow(args)
    elif args.mode == "web-to-workflow":
        cmd_web_to_workflow(args)
    elif args.mode == "webpage-to-workflow":
        cmd_webpage_to_workflow(args)


if __name__ == "__main__":
    main()
